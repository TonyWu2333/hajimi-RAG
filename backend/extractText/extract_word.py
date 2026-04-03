"""
extract_word.py
提取 Word 文档（.docx / .doc）中的纯文本。

依赖:
    pip install python-docx
    .doc 格式还需要安装 LibreOffice（用于格式转换）

用法:
    from extract_word import extract_word
    text = extract_word("path/to/file.docx")
"""

import subprocess
import tempfile
from pathlib import Path


def _convert_doc_to_docx(doc_path: Path) -> Path:
    """使用 LibreOffice 将旧版 .doc 转换为 .docx，返回临时文件路径。"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to", "docx",
                "--outdir", tmp_dir,
                str(doc_path),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice 转换失败: {result.stderr.strip()}"
            )
        converted = Path(tmp_dir) / (doc_path.stem + ".docx")
        if not converted.exists():
            raise RuntimeError("转换后文件未找到，请确认 LibreOffice 已安装。")
        # 复制到系统临时目录，避免 TemporaryDirectory 清理
        import shutil
        dest = Path(tempfile.mktemp(suffix=".docx"))
        shutil.copy2(converted, dest)
        return dest


def extract_word(file_path: str) -> str:
    """
    提取 Word 文档的文本内容。

    支持 .docx（原生）和 .doc（需要 LibreOffice）。
    提取段落文本和表格单元格文本。

    Args:
        file_path: Word 文件路径（.docx 或 .doc）。

    Returns:
        提取的纯文本字符串。

    Raises:
        FileNotFoundError: 文件不存在。
        ImportError: 未安装 python-docx。
        ValueError: 文件格式不支持。
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请先安装依赖: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    if suffix not in (".docx", ".doc"):
        raise ValueError(f"不支持的文件格式: {suffix}，仅支持 .docx / .doc")

    tmp_path = None
    try:
        if suffix == ".doc":
            tmp_path = _convert_doc_to_docx(path)
            docx_path = tmp_path
        else:
            docx_path = path

        doc = Document(str(docx_path))
        parts = []

        # 段落
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells]
                line = "\t".join(row_texts)
                if line.strip():
                    parts.append(line)

        return "\n".join(parts)

    finally:
        if tmp_path and tmp_path.exists():
            tmp_path.unlink(missing_ok=True)


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("用法: python extract_word.py <file.docx>")
        sys.exit(1)
    print(extract_word(sys.argv[1]))